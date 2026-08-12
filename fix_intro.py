import docx

doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')

# 1. Clear everything between '1. Introduction' and '2. Method'
in_intro = False
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('1.') and 'Introduction' in text:
        in_intro = True
        continue
    
    if in_intro:
        if text.startswith('2.') and 'Method' in text:
            in_intro = False
            continue
        # Clear paragraph
        p.text = ""

# Remove empty paragraphs
for p in doc.paragraphs:
    if p.text == "":
        p._element.getparent().remove(p._element)

# 2. Insert new Introduction before '2. Method'
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('2.') and 'Method' in text:
        intro_text = [
            "1.1 Research Background",
            "The Free Nutritious Meal Program (MBG) is one of Indonesia's national strategic programs aimed at improving students' nutritional quality while supporting the development of healthy and productive human resources [1]. Through this program, nutritious meals are distributed daily to students across Indonesia. While the program has significant social benefits, its implementation faces operational challenges related to monitoring, nutritional evaluation, and food waste management [2]. According to waste management communities in Surabaya, approximately 40 to 50 kg of food waste can be generated each day from a single school [3]. This highlights the urgent need for a digital monitoring system capable of integrating food distribution tracking and sustainable waste management.",
            "To address this gap, an integrated web-based platform named SmartMBG is proposed. For such a platform to be adopted successfully by its diverse users—which include teachers, school administrators, nutrition service units (SPPG), and organic waste partners—the system must possess a highly intuitive and accessible user interface (UI) and user experience (UX) [4]. Poor interface design often leads to user resistance, data entry errors, and overall system failure in public service applications [5]. Therefore, designing an effective UI/UX is a critical phase in the development of the SmartMBG platform.",
            "1.2 Related Works and Research Gap",
            "Previous studies have emphasized the importance of user-centered approaches in developing monitoring dashboards and public sector applications [6]. However, few studies have specifically explored the UI/UX design for school nutrition and food waste monitoring systems using a structured methodology. Most existing systems still rely on manual data entry or possess complex interfaces that burden the users [7].",
            "This study aims to fill that research gap by applying the Design Thinking methodology to design the UI/UX of the SmartMBG platform. Design Thinking is a human-centered approach to innovation that integrates the needs of people, the possibilities of technology, and the requirements for project success [8]. By utilizing the Empathize, Define, Ideate, Prototype, and Test phases, this research ensures that the resulting digital interface is closely aligned with the actual needs and technical proficiencies of the end-users.",
            "1.3 Research Objectives",
            "The primary objective of this research is to design and evaluate the User Interface (UI) and User Experience (UX) of the SmartMBG platform using the Design Thinking methodology. The specific goals include identifying user pain points through empathy research, creating wireframes and high-fidelity prototypes in Figma, and conducting usability testing to ensure the design meets user requirements."
        ]
        
        for t in intro_text:
            p.insert_paragraph_before(t)
        break

doc.save('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
print("Introduction replaced successfully.")
