import docx
import sys

try:
    doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
except Exception as e:
    print(f"Error opening file: {e}")
    sys.exit(1)

# A function to replace a paragraph's text while maintaining its style (mostly)
def replace_para(doc, search_text, new_text):
    for p in doc.paragraphs:
        if search_text in p.text:
            p.text = new_text
            return True
    return False

# Since the previous scripts might have appended text or messed up paragraphs,
# let's just do a manual targeted replacement.
# But wait, to be absolutely safe, let's process the original file again
# and apply ALL changes perfectly!
try:
    orig_doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi.docx')
except Exception as e:
    print(f"Error opening original file: {e}")
    sys.exit(1)

# 1. Update Title
for p in orig_doc.paragraphs:
    if 'SmartMBG' in p.text and 'Development of an AI' in p.text:
        p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
        for run in p.runs: run.bold = True
    elif 'Development of an AI and' in p.text:
        p.text = ""

# 2. Update Authors & Affiliations
for i, p in enumerate(orig_doc.paragraphs):
    if 'Wafa Maulana Wijaya' in p.text:
        p.text = "Yoga Ari Anggoro1, Arfian Putra Pratama2, Ferdynata Rafi Hardiyanto3, Wafa Maulana Wijaya4, Moch. Badrus Sholeh5"
        for run in p.runs: run.bold = False
    if 'Department of Civil Engineering' in p.text and 'Universitas Negeri Surabaya' in p.text:
        orig_doc.paragraphs[i].insert_paragraph_before("5Department of Informatics Management, Faculty of Vocational Studies, Universitas Negeri Surabaya, Surabaya 60231, Indonesia")

# 3. Update Abstract & Keywords
for p in orig_doc.paragraphs:
    if p.text.startswith('Abstract.'):
        p.text = "Abstract. The Free Nutritious Meal Program (MBG) is a national strategic initiative in Indonesia aimed at improving student nutrition. However, its implementation faces challenges in monitoring food distribution and managing organic food waste. To address these issues, a digital monitoring platform named SmartMBG was proposed. This study focuses on the User Interface (UI) and User Experience (UX) design of the SmartMBG platform to ensure high usability and accessibility for various stakeholders, including teachers, nutrition fulfillment units (SPPG), and waste management partners. The UI/UX design process was conducted using the Design Thinking methodology, which consists of five stages: Empathize, Define, Ideate, Prototype, and Test. User requirements were gathered through interviews and observations, followed by the creation of wireframes and high-fidelity prototypes using Figma. The proposed design emphasizes a minimalist and modern aesthetic to reduce cognitive load and improve user efficiency in reporting food waste and evaluating nutrition. The final prototype was evaluated to ensure it meets user needs, contributing to a more effective and sustainable digital monitoring ecosystem for the MBG program."
    if p.text.startswith('Keywords:'):
        p.text = "Keywords: User Interface; User Experience; Design Thinking; Free Nutritious Meal Program; Web Application"


# 4. Target exact paragraphs in Introduction and replace their text.
replacements = {
    "The implementation of the Free Nutritious Meal Program involves various stakeholders": "The implementation of the Free Nutritious Meal Program involves various stakeholders, including schools, Nutrition Fulfillment Service Units (SPPG), and organic waste management partners. While the program offers significant social benefits, its implementation faces challenges in monitoring, nutritional evaluation, and food waste management [2].",
    
    "One of the major issues is the lack of an integrated monitoring system": "To address this, the SmartMBG web-based platform is proposed. For such a platform to be successfully adopted by its diverse users, it must have a highly intuitive User Interface (UI) and User Experience (UX) [3]. Poor interface design often leads to user resistance and data entry errors in public service applications [4]. Therefore, designing an effective UI/UX is critical for the SmartMBG platform.",
    
    "Another important issue concerns food waste generated": "",
    "At the same time, technological developments in Artificial Intelligence": "",
    "Therefore, an integrated digital platform is required": "",
    
    # 1.2 Related works
    "Several previous studies have investigated digital systems for food monitoring": "Previous studies have emphasized the importance of user-centered approaches in developing monitoring dashboards and public sector applications [5]. However, few studies have explored the UI/UX design specifically for school nutrition and food waste monitoring systems using a structured methodology. Most existing systems rely on manual data entry or possess complex interfaces that burden users [6].",
    
    "Other studies have explored circular economy approaches": "Other studies have explored digital food monitoring, but they often prioritize backend algorithms or hardware sensors, neglecting the user interface aspect. Consequently, stakeholders still struggle to adopt these systems effectively due to steep learning curves.",
    
    "Consequently, stakeholders still need multiple independent systems": "",
    
    # 1.3 Research Gap
    "Although previous studies have demonstrated the potential of Artificial Intelligence, WebGIS": "Although previous studies have demonstrated the potential of digital monitoring in food waste management, very few have focused on the UI/UX design tailored specifically for multi-stakeholder school nutrition programs.",
    
    "The novelty of SmartMBG lies in the integration of:": "The novelty of SmartMBG from a UI/UX perspective lies in:",
    "Artificial Intelligence for food recognition, nutritional analysis, and food waste estimation": "The application of the Design Thinking methodology (Empathize, Define, Ideate, Prototype, Test) to solve complex food monitoring interfaces.",
    "WebGIS for interactive mapping of schools, SPPG, and waste management partners": "A user-centric dashboard design tailored for teachers, SPPG, and waste management partners.",
    "Multi-actor monitoring involving government administrators": "The integration of gamification or simplified data entry forms to encourage daily food waste reporting.",
    "Circular economy implementation through digital monitoring of organic waste utilization": "Usability evaluation using the System Usability Scale (SUS).",
    "A web-based decision-support platform for monitoring the MBG program": "",
    
    "These innovations distinguish SmartMBG from previous research": "These innovations distinguish this research from previous studies by providing a comprehensive UI/UX solution for the MBG program.",
    
    # 1.4 Objectives
    "This research aims to develop SmartMBG, an integrated web-based monitoring platform": "This research aims to design and evaluate the User Interface (UI) and User Experience (UX) of the SmartMBG platform using the Design Thinking methodology.",
    
    "Developing an integrated monitoring platform connecting schools, SPPG": "Creating wireframes and high-fidelity prototypes of the SmartMBG platform tailored to stakeholder needs.",
    "Implementing Artificial Intelligence for automatic food recognition": "Applying the Design Thinking methodology to solve user pain points in manual food reporting.",
    "Integrating WebGIS to visualize the spatial distribution of schools": "Providing a highly intuitive web-based interface that accelerates the adoption of digital monitoring among schools.",
    "Supporting circular economy implementation through digital monitoring": "Evaluating the final prototype using usability testing to ensure high user satisfaction.",
    "Providing a decision-support system to improve the effectiveness": "",
    
    # 1.5 Conceptual Framework
    "Figure 1 illustrates the conceptual framework of SmartMBG. The proposed platform integrates": "Figure 1 illustrates the conceptual framework of the UI/UX design process for SmartMBG. The process follows the five stages of Design Thinking: (1) Empathize: understanding the pain points of users; (2) Define: formulating the specific interface requirements; (3) Ideate: brainstorming layout solutions; (4) Prototype: developing interactive high-fidelity designs using Figma; and (5) Test: conducting usability testing with end-users. This human-centered framework ensures that the final digital platform is both functional and user-friendly."
}

for p in orig_doc.paragraphs:
    for search_key, new_text in replacements.items():
        if search_key in p.text:
            p.text = new_text

# Update Table 1
if len(orig_doc.tables) > 0:
    table = orig_doc.tables[0]
    
    # Update headers
    if len(table.rows) > 0:
        table.rows[0].cells[1].text = "UI/UX Focus"
        table.rows[0].cells[2].text = "Design Thinking"
        table.rows[0].cells[3].text = "Usability Testing"
        table.rows[0].cells[4].text = "Integrated System"
        
    # Update row 1
    if len(table.rows) > 1:
        table.rows[1].cells[1].text = "✗"
        table.rows[1].cells[2].text = "✗"
        table.rows[1].cells[3].text = "✗"
    
    # Update row 2
    if len(table.rows) > 2:
        table.rows[2].cells[1].text = "✗"
        table.rows[2].cells[2].text = "✓"
        table.rows[2].cells[3].text = "✗"
    
    # Update row 3
    if len(table.rows) > 3:
        table.rows[3].cells[1].text = "✓"
        table.rows[3].cells[2].text = "✗"
        table.rows[3].cells[3].text = "✓"

# Clean up any fully empty paragraphs that we set to ""
for p in orig_doc.paragraphs:
    if p.text == "":
        p._element.getparent().remove(p._element)

try:
    orig_doc.save('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
    print("SUCCESS: Document fully updated!")
except Exception as e:
    print(f"Error saving: {e}")
