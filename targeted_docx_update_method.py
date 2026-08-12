import docx
import sys

try:
    doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
except Exception as e:
    print(f"Error opening file: {e}")
    sys.exit(1)

replacements = {
    # 2. Method
    "This research employed the Research and Development (R&D) approach to design and develop SmartMBG": "This research employed the Design Thinking methodology to design and develop the User Interface (UI) and User Experience (UX) of SmartMBG, an integrated web-based platform for monitoring the implementation of Indonesia's Free Nutritious Meal Program (MBG). The Design Thinking approach was chosen because it provides a solution-based approach to solving complex problems by understanding the human needs involved.",
    
    "The proposed system integrates Artificial Intelligence (AI), WebGIS, and circular economy concepts": "",
    
    "The research consisted of several sequential stages, including problem identification, literature review, system requirements analysis, system design, system development, implementation of Artificial Intelligence and WebGIS, system testing and evaluation, and documentation of research outputs.": "The research consisted of five sequential stages: Empathize, Define, Ideate, Prototype, and Test.",
    
    "The system was developed using Laravel as the backend framework, React.js as the frontend framework, PostgreSQL as the database management system, Leaflet.js for WebGIS visualization, and YOLO-based Computer Vision": "The system's interface was designed using Figma as the primary prototyping tool. Functional and usability verification of the proposed design was conducted using the System Usability Scale (SUS), while user acceptance was evaluated based on direct feedback from stakeholders.",
    
    "Functional verification of the developed system was conducted using Black Box Testing": "",
    
    # 2.1 Research Design
    "The research adopted the Research and Development (R&D) methodology because the primary objective was to develop an integrated software platform rather than testing a specific hypothesis.": "The research adopted the Design Thinking methodology because the primary objective was to develop a user-centered interface rather than testing a specific hypothesis.",
    
    "The development process followed several iterative stages beginning with problem identification and ending with system evaluation.": "The development process followed several iterative stages beginning with empathy and ending with user testing.",
    
    "Initially, the existing implementation of the MBG program was analyzed to identify operational problems": "Initially, the existing implementation of the MBG program was analyzed to empathize with users and identify operational problems",
    
    "Subsequently, literature related to Artificial Intelligence, WebGIS, food waste management, and circular economy was reviewed to establish the theoretical foundation for system development.": "Subsequently, user personas and journey maps were created during the define and ideate phases to establish the foundation for the interface design.",
    
    "Based on the identified requirements, SmartMBG was designed using a modular architecture consisting of Artificial Intelligence, WebGIS, and monitoring dashboard components.": "Based on the identified requirements, SmartMBG's UI/UX was designed using Figma.",
    
    "The system architecture was then implemented into a web-based application, followed by functional testing and system evaluation.": "The wireframes and high-fidelity prototypes were then subjected to usability testing and system evaluation with target users.",
    
    # 2.2 System Architecture -> Empathize and Define Phases
    "2.2 System Architecture": "2.2 Empathize and Define Phases",
    
    "The SmartMBG platform adopts a three-layer architecture consisting of the Presentation Layer, Application Layer, and Data Layer.": "The Empathize phase aimed to gain a deep understanding of the target users—teachers, school administrators, SPPG, and waste management partners. Through in-depth interviews and field observations in Surabaya, the researchers gathered insights into users' daily workflows, pain points, and technical proficiencies.",
    
    "This layered architecture is designed to improve system scalability, maintainability, and communication among stakeholders while supporting the integration of Artificial Intelligence (AI) and WebGIS technologies.": "",
    
    "The Presentation Layer provides web-based user interfaces for four categories of users": "Following empathy mapping, the Define phase synthesized these findings to formulate clear problem statements. It became evident that stakeholders required a simplified dashboard, straightforward navigation, and quick data entry forms. Poor interface design in previous tools had led to user resistance, making an intuitive layout critical.",
    
    "Through this interface, users can perform activities such as monitoring food distribution": "",
    "The Application Layer serves as the core of the SmartMBG platform by managing business logic using the Laravel framework.": "",
    "This layer also provides RESTful APIs to facilitate communication between the frontend application and backend services.": "",
    "Furthermore, it integrates Artificial Intelligence (AI) for food recognition": "",
    "The Data Layer utilizes PostgreSQL as the centralized database management system": "",
    "Centralized data storage ensures data consistency, supports efficient information retrieval": "",
    "Figure 3 illustrates the overall architecture of the SmartMBG platform.": "",
    "The architecture demonstrates how users interact with the web-based interface in the Presentation Layer": "",
    "This layered architecture enables seamless communication among all stakeholders": "",
    
    "Fig. 3. SmartMBG System Architecture": "Fig. 3. User Journey and Flow Architecture",
    
    # 2.3 Artificial Intelligence Module -> Ideate Phase
    "2.3 Artificial Intelligence Module": "2.3 Ideate Phase",
    "The Artificial Intelligence module is responsible for automatically identifying food items and estimating nutritional content from food images uploaded by users.": "The Ideate phase focused on generating a wide range of design solutions to address the defined problems. Brainstorming sessions were conducted to sketch various layout configurations, navigation menus, and data visualization formats.",
    "The AI module applies YOLO (You Only Look Once) object detection to recognize food objects.": "Key design decisions included adopting a minimalist and modern aesthetic, utilizing a distinct color palette to differentiate user roles, and incorporating clear calls-to-action (CTAs) for uploading food waste reports and viewing nutritional analytics.",
    "After food recognition, nutritional values are estimated by matching identified food items with nutritional databases.": "",
    "The AI module also estimates food waste by analyzing leftover food captured through uploaded images.": "",
    "The estimated nutritional value can be represented as": "",
    "where": "",
    "N = estimated nutritional value,": "",
    "Wᵢ = estimated food weight,": "",
    "Cᵢ = nutritional composition,": "",
    "n = number of detected food items.": "",
    
    # 2.4 WebGIS Module -> Prototype Phase
    "2.4 WebGIS Module": "2.4 Prototype Phase",
    "The WebGIS module provides geographic visualization of the MBG ecosystem.": "The Prototype phase transformed the selected ideas into tangible, interactive representations of the SmartMBG platform.",
    "The module displays:": "The design process began with low-fidelity wireframes to establish the structural layout. Subsequently, high-fidelity prototypes were developed using Figma. These prototypes simulated the actual web experience, complete with interactive buttons, smooth page transitions, and responsive design elements tailored for both desktop and mobile views.",
    "school locations,": "",
    "SPPG locations,": "",
    "waste management partner locations,": "",
    "distribution coverage,": "",
    "monitoring status.": "",
    "Interactive maps are implemented using Leaflet.js, while geographic coordinates are stored in PostgreSQL": "",
    
    # 2.5 System Testing -> Usability Testing
    "2.5 System Testing": "2.5 Usability Testing",
    "System testing was performed using the Black Box Testing approach.": "Usability testing was performed to evaluate how well the target users could navigate and interact with the SmartMBG prototype.",
    "Each functional module was tested according to predefined scenarios to verify whether the expected outputs were generated correctly.": "Each stakeholder was given specific tasks to complete, such as logging in, submitting a food waste report, or viewing the monitoring map. The evaluation utilized the System Usability Scale (SUS), a widely accepted metric for measuring user satisfaction.",
    "The tested modules include": "The success rate and usability score were calculated using the standard SUS formula:",
    "Login": "",
    "Dashboard": "",
    "User Management": "",
    "School Management": "",
    "SPPG Management": "",
    "Food Waste Reporting": "",
    "AI Analysis": "",
    "WebGIS": "",
    "Reporting": "",
    "Notification": "",
    "The success rate was calculated using": "",
    
    # 2.6 Data Analysis
    "The collected data were analyzed descriptively.": "The collected data from the usability testing were analyzed quantitatively and qualitatively.",
    "Functional test results were summarized using percentage-based analysis, while the AI model performance was evaluated based on prediction accuracy.": "Quantitative analysis involved calculating the final SUS scores to determine the overall usability rating (e.g., Acceptable, Marginal, or Not Acceptable). Qualitative data obtained from user feedback and direct observation during testing were analyzed descriptively to identify areas for interface refinement.",
    "Spatial data generated by WebGIS were analyzed to observe the distribution of schools, SPPG, and waste management partners.": "",
    "Finally, the developed platform was evaluated based on functionality, usability, and system integration to determine whether SmartMBG successfully addressed the identified research problems.": "Finally, the designed prototype was iteratively improved based on these evaluation results to ensure that SmartMBG successfully addressed the stakeholders' needs."
}

for p in doc.paragraphs:
    # Some texts might be split into runs, so doing a simple string replace on p.text 
    # directly modifies it properly via docx API. But p.text = ... overrides all formatting.
    # Since these are body paragraphs, overriding formatting is usually fine.
    original_text = p.text
    new_text = original_text
    
    # To handle exact matching even with minor spaces
    for search_key, replace_text in replacements.items():
        if search_key in new_text:
            new_text = new_text.replace(search_key, replace_text)
            
    if new_text != original_text:
        p.text = new_text

# Clean up any fully empty paragraphs that we set to ""
for p in doc.paragraphs:
    if p.text.strip() == "" and len(p.runs) == 0:
        pass # Better not to mess with empty paragraphs as they might be used for spacing or images

# We specifically set a lot of list items to "" (like • school locations, etc.)
for p in doc.paragraphs:
    if p.text == "•\t" or p.text == "•":
        p.text = ""
        
# Custom cleanup for bullets
for p in doc.paragraphs:
    if p.text in ['school locations,', 'SPPG locations,', 'waste management partner locations,', 'distribution coverage,', 'monitoring status.', 'Login', 'Dashboard', 'User Management', 'School Management', 'SPPG Management', 'Food Waste Reporting', 'AI Analysis', 'WebGIS', 'Reporting', 'Notification']:
        p.text = ""

try:
    doc.save('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
    print("SUCCESS: Chapter 2 Method fully updated!")
except Exception as e:
    print(f"Error saving: {e}")
