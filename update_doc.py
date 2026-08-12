import docx

# Open the ORIGINAL document to start fresh
doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi.docx')

# 1. Update Title
for p in doc.paragraphs:
    if 'SmartMBG' in p.text and 'Development of an AI' in p.text:
        p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
        for run in p.runs: run.bold = True
        break
    elif 'Development of an AI and' in p.text:
        p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
        for run in p.runs: run.bold = True
        break

# 2. Update Authors
for p in doc.paragraphs:
    if 'Wafa Maulana Wijaya' in p.text:
        p.text = "Yoga Ari Anggoro1, Arfian Putra Pratama2, Ferdynata Rafi Hardiyanto3, Wafa Maulana Wijaya4, Moch. Badrus Sholeh5"
        for run in p.runs:
            run.bold = False
        break

# 3. Update Affiliations
for i, p in enumerate(doc.paragraphs):
    if 'Department of Civil Engineering' in p.text and 'Universitas Negeri Surabaya' in p.text:
        # Insert 5th affiliation after this paragraph
        new_p = doc.paragraphs[i].insert_paragraph_before("5Department of Informatics Management, Faculty of Vocational Studies, Universitas Negeri Surabaya, Surabaya 60231, Indonesia")
        break

# 4. Update Abstract
for p in doc.paragraphs:
    if p.text.startswith('Abstract.'):
        p.text = "Abstract. The Free Nutritious Meal Program (MBG) is a national strategic initiative in Indonesia aimed at improving student nutrition. However, its implementation faces challenges in monitoring food distribution and managing organic food waste. To address these issues, a digital monitoring platform named SmartMBG was proposed. This study focuses on the User Interface (UI) and User Experience (UX) design of the SmartMBG platform to ensure high usability and accessibility for various stakeholders, including teachers, nutrition fulfillment units (SPPG), and waste management partners. The UI/UX design process was conducted using the Design Thinking methodology, which consists of five stages: Empathize, Define, Ideate, Prototype, and Test. User requirements were gathered through interviews and observations, followed by the creation of wireframes and high-fidelity prototypes using Figma. The proposed design emphasizes a minimalist and modern aesthetic to reduce cognitive load and improve user efficiency in reporting food waste and evaluating nutrition. The final prototype was evaluated to ensure it meets user needs, contributing to a more effective and sustainable digital monitoring ecosystem for the MBG program."
        break

# 5. Update Keywords
for p in doc.paragraphs:
    if p.text.startswith('Keywords:'):
        p.text = "Keywords: User Interface; User Experience; Design Thinking; Free Nutritious Meal Program; Web Application"
        break

# 6. Update Introduction
in_intro = False
for p in doc.paragraphs:
    if p.text.strip() == '1. Introduction' or p.text.strip() == 'Introduction':
        in_intro = True
        continue
    
    if in_intro:
        if '2. Method' in p.text or p.text.strip() == '2. Method':
            in_intro = False
            continue
        # Clear all text in introduction section to rewrite it
        p.text = ""

# Now we will insert the new Introduction
for i, p in enumerate(doc.paragraphs):
    if '2. Method' in p.text or p.text.strip() == '2. Method':
        intro_text = [
            "1.1 Research Background",
            "The Free Nutritious Meal Program (MBG) is one of Indonesia's national strategic programs aimed at improving students' nutritional quality while supporting the development of healthy and productive human resources [1]. Through this program, nutritious meals are distributed daily to students across Indonesia. While the program has significant social benefits, its implementation faces operational challenges related to monitoring, nutritional evaluation, and food waste management [2]. According to waste management communities in Surabaya, approximately 40 to 50 kg of food waste can be generated each day from a single school [3]. This highlights the urgent need for a digital monitoring system capable of integrating food distribution tracking and sustainable waste management.",
            "To address this gap, an integrated web-based platform named SmartMBG is proposed. For such a platform to be adopted successfully by its diverse users—which include teachers, school administrators, nutrition service units (SPPG), and organic waste partners—the system must possess a highly intuitive and accessible user interface (UI) and user experience (UX) [4]. Poor interface design often leads to user resistance, data entry errors, and overall system failure in public service applications [5]. Therefore, designing an effective UI/UX is a critical phase in the development of the SmartMBG platform.",
            "1.2 Related Works and Research Gap",
            "Previous studies have emphasized the importance of user-centered approaches in developing monitoring dashboards and public sector applications [6]. However, few studies have specifically explored the UI/UX design for school nutrition and food waste monitoring systems using a structured methodology. Most existing systems still rely on manual data entry or possess complex interfaces that burden the users [7].",
            "This study aims to fill that research gap by applying the Design Thinking methodology to design the UI/UX of the SmartMBG platform. Design Thinking is a human-centered approach to innovation that integrates the needs of people, the possibilities of technology, and the requirements for project success [8]. By utilizing the Empathize, Define, Ideate, Prototype, and Test phases, this research ensures that the resulting digital interface is closely aligned with the actual needs and technical proficiencies of the end-users.",
            "1.3 Research Objectives",
            "The primary objective of this research is to design and evaluate the User Interface (UI) and User Experience (UX) of the SmartMBG platform using the Design Thinking methodology. The specific goals include identifying user pain points through empathy research, creating wireframes and high-fidelity prototypes in Figma, and conducting usability testing to ensure the design meets user requirements."
        ]
        
        # Insert paragraphs normally (forward iteration) before "2. Method"
        for text in intro_text:
            p.insert_paragraph_before(text)
        break

# Clean up empty paragraphs
for p in doc.paragraphs:
    if p.text == "":
        p._element.getparent().remove(p._element)

doc.save('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
print("Document updated successfully.")
