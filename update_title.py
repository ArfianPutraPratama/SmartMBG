import docx

doc = docx.Document('SmartMBG_Yoga Ari Anggoro_Icasvi.docx')

title_found = False
for p in doc.paragraphs:
    # Check if this paragraph contains part of the original title
    if 'SmartMBG' in p.text and 'Development of an AI' in p.text:
        # Clear the old text and set the new one
        p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
        # Since it's a title, let's make it bold and centered if it was
        for run in p.runs:
            run.bold = True
        title_found = True
        break
    elif 'Development of an AI and' in p.text:
        p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
        for run in p.runs:
            run.bold = True
        title_found = True
        break

# If the title was split across multiple runs or lines, let's just replace the first paragraph if it contains SmartMBG and Development
if not title_found:
    for p in doc.paragraphs:
        if 'SmartMBG' in p.text:
            print("Found paragraph with SmartMBG:", p.text)
            p.text = "Perancangan User Interface dan User Experience (UI/UX) pada Platform Pemantauan Program Makan Bergizi (SmartMBG) Menggunakan Metode Design Thinking"
            for run in p.runs:
                run.bold = True
            break

doc.save('SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx')
print("Saved as SmartMBG_Yoga Ari Anggoro_Icasvi_UIUX.docx")
